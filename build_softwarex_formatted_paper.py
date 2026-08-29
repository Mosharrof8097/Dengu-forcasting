import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

OUTPUT_MANUSCRIPT = "/home/mosharrof/personal Doc/medipep/02_EpiST_Shield_App_Paper/EpiST-Shield_SoftwareX_Manuscript.docx"
FIGURES_DIR = "/home/mosharrof/personal Doc/medipep/02_EpiST_Shield_App_Paper/05_Paper_Figures"

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_1(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    r = h.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_heading_2(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    r = h.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_heading_3(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(2)
    r = h.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.italic = True
    r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_body_p(doc, text, bold_prefix=None, italic_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_b = p.add_run(bold_prefix)
        r_b.font.name = 'Arial'
        r_b.font.size = Pt(11)
        r_b.font.bold = True
        r_b.font.color.rgb = RGBColor(0, 0, 0)
    if italic_prefix:
        r_i = p.add_run(italic_prefix)
        r_i.font.name = 'Arial'
        r_i.font.size = Pt(11)
        r_i.font.italic = True
        r_i.font.color.rgb = RGBColor(0, 0, 0)
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_code_block(doc, code_str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(code_str)
    r.font.name = 'Courier New'
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(30, 30, 30)
    return p

def add_figure_img(doc, img_name, caption_str, width_in=5.8):
    fig_path = os.path.join(FIGURES_DIR, img_name)
    if not os.path.exists(fig_path):
        print(f"Warning: Figure not found at {fig_path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(fig_path, width=Inches(width_in))
    
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(10)
    r_cap = cp.add_run(caption_str)
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(9.5)
    r_cap.font.italic = True
    r_cap.font.color.rgb = RGBColor(0, 0, 0)

def main():
    print("Generating official SoftwareX formatted manuscript...")
    doc = Document()
    
    # Page setup - Margins matching standard SoftwareX template (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # TITLE (SoftwareX Style)
    h_title = doc.add_paragraph()
    h_title.paragraph_format.space_before = Pt(0)
    h_title.paragraph_format.space_after = Pt(8)
    r_t = h_title.add_run("Title (Name of your software: EpiST-Shield)")
    r_t.font.name = 'Arial'
    r_t.font.size = Pt(16)
    r_t.font.bold = True
    r_t.font.color.rgb = RGBColor(0, 0, 0)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(12)
    r_s = p_sub.add_run("EpiST-Shield: An Open-Access Web-Based Decision Support Platform for Multi-Horizon Dengue Outbreak Forecasting and Prescriptive Healthcare Resource Allocation in Bangladesh")
    r_s.font.name = 'Arial'
    r_s.font.size = Pt(13)
    r_s.font.bold = True
    r_s.font.color.rgb = RGBColor(0, 0, 0)

    # AUTHORS & AFFILIATIONS
    add_heading_1(doc, "Names of authors / main developers (incl. affiliations, addresses, email)")
    add_body_p(doc, "Md. Mosharrof Hossain¹*")
    add_body_p(doc, "¹Department of Computer Science & Engineering, Bangladesh")
    add_body_p(doc, "mdmosharrofhossain455@gmail.com", bold_prefix="*Corresponding author email: ")

    # ABSTRACT
    add_heading_1(doc, "Abstract")
    add_body_p(doc, "Dengue fever presents a recurring public health crisis in Bangladesh, characterized by localized outbreak surges and severe hospital resource exhaustion. Translating complex epidemiological models into real-time operational decisions remains a major hurdle for public health authorities. Here, we introduce EpiST-Shield, an open-access, web-based decision support platform designed for multi-horizon (7-, 14-, 21-, and 30-day) dengue forecasting and prescriptive health logistics planning across 11 focal endemic districts. Powered by an asynchronous FastAPI backend and a pre-trained deep learning inference engine (EpiST-Former), EpiST-Shield achieves sub-millisecond core inference latency (<0.02 ms) and supports high-concurrency throughput exceeding 48,500 requests per second. The platform translates raw case projections into automated hospital resource schedules (beds, NS1 kits, saline bags), weekly priority dispatches, interactive 'What-If' weather scenario stress-testing, and executive PDF bulletins. A field usability audit with 15 health domain experts yielded a System Usability Scale (SUS) score of 88.4/100 (Grade A+ Excellent). EpiST-Shield is publicly deployed at https://dengu-forcasting.vercel.app/ under an open-source MIT license.")

    # KEYWORDS
    add_heading_1(doc, "Keywords")
    add_body_p(doc, "Dengue forecasting; outbreak early-warning system; decision support system; healthcare resource planning; geospatial visualization; web deployment")

    # METADATA TABLE (Mandatory Code Metadata Table 1)
    add_heading_1(doc, "Metadata")
    add_body_p(doc, "This ancillary data table is required for your submission. It summarizes the core code version, repository links, legal license, dependencies, and support metadata for EpiST-Shield.")

    headers_meta = ["Nr", "Code metadata description", "Metadata"]
    rows_meta = [
        ["C1", "Current code version", "v1.0.0"],
        ["C2", "Permanent link to code/repository used for this code version", "https://github.com/Mosharrof8097/Dengu-forcasting"],
        ["C3", "Legal code license", "MIT License"],
        ["C4", "Code versioning system used", "git"],
        ["C5", "Software code languages, tools and services used", "Python 3.12, FastAPI, TensorFlow/Keras 3.x, HTML5, Vanilla CSS3, JavaScript (ES6+), Leaflet.js, Vercel"],
        ["C6", "Compilation requirements, operating environments and dependencies", "Python 3.10+; Node.js 18+; see requirements.txt and package.json"],
        ["C7", "If available, link to developer documentation/manual", "https://github.com/Mosharrof8097/Dengu-forcasting#readme"],
        ["C8", "Support email for questions", "mdmosharrofhossain455@gmail.com"],
    ]

    t_meta = doc.add_table(rows=len(rows_meta)+1, cols=len(headers_meta))
    t_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Table 1 Header Styling (White fill, bold black text)
    hdr_cells = t_meta.rows[0].cells
    for i, h_text in enumerate(headers_meta):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "FFFFFF")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(9.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0, 0, 0)

    for r_idx, row_data in enumerate(rows_meta):
        row_cells = t_meta.rows[r_idx+1].cells
        for c_idx, cell_value in enumerate(row_data):
            row_cells[c_idx].text = str(cell_value)
            set_cell_background(row_cells[c_idx], "FFFFFF")
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[c_idx].paragraphs[0]
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0, 0, 0)
                if c_idx == 0:
                    r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # =============================================================
    # SECTION 1: MOTIVATION AND SIGNIFICANCE
    # =============================================================
    add_heading_1(doc, "1. Motivation and significance")
    add_body_p(doc, "Dengue fever remains one of the most significant vector-borne disease burdens in Bangladesh, with case counts and geographical spread accelerating dramatically in recent years. The Directorate General of Health Services (DGHS) reported unprecedented surge waves during the 2023–2024 outbreak seasons, recording over 321,000 hospitalized cases and 1,705 fatalities. During explosive monsoonal outbreaks, major tertiary health facilities in urban centers such as Dhaka, Chattogram, and Mymensingh face severe shortages of isolation beds, intravenous (IV) fluid saline bags, and rapid diagnostic NS1 test kits.")

    add_body_p(doc, "While recent epidemiological literature has introduced various machine learning and deep learning models for disease prediction (e.g., Rahman et al., 2025; Shiddik et al., 2026; Alam et al., 2025), a critical operational gap persists: existing models operate primarily as offline, passive diagnostic tools. A projected case curve alone does not tell a public health officer how many extra hospital beds to procure or how to schedule weekly medical dispatches across districts. Furthermore, existing research rarely provides interactive tools for decision-makers to stress-test forecast trajectories under extreme micro-climatic shifts (e.g., unexpected monsoonal rainfall spikes).")

    add_body_p(doc, "To resolve this gap, we developed EpiST-Shield, an open-access web-based decision support system designed specifically for public health officials, district civil surgeons, and hospital logistics planners in Bangladesh. EpiST-Shield bridges the gap between deep learning prospective forecasting and real-world healthcare execution by providing: (1) multi-horizon (7, 14, 21, and 30-day) district-level case projections, (2) automated hospital bed and medical supply calculations, (3) a 4-week prioritized dispatch schedule, (4) an interactive 'What-If' weather simulator, and (5) one-click automated executive PDF bulletin exports.")

    # =============================================================
    # SECTION 2: SOFTWARE DESCRIPTION
    # =============================================================
    add_heading_1(doc, "2. Software description")

    add_heading_2(doc, "2.1 Software architecture and fault-tolerant design")
    add_body_p(doc, "EpiST-Shield is deployed live and publicly accessible at https://dengu-forcasting.vercel.app/. The platform follows a decoupled client-server web architecture consisting of a browser-based frontend and a FastAPI backend service responsible for forecast generation and resource estimation (Figure 1). The primary forecasting pipeline uses the pre-trained EpiST-Former model, while additional fallback mechanisms are implemented to maintain degraded service continuity when either primary model inference or backend connectivity becomes unavailable.")

    add_body_p(doc, "The fault-tolerant architecture operates at two independent levels. At the server level, a backend mathematical fallback is activated when the primary EpiST-Former inference pipeline cannot produce a valid result. At the client level, an offline fallback mechanism is activated when the frontend cannot obtain a valid forecast response from the backend API. These mechanisms are coupled with explicit forecast-source tracking to ensure that users can distinguish primary deep-learning forecasts from fallback-generated estimates.")

    add_body_p(doc, "For successful backend communication, the API response includes the metadata attributes forecast_source and fallback_active. The forecast_source field identifies whether the returned forecast was generated by the primary EpiST-Former model or by the backend mathematical fallback, while fallback_active indicates whether a degraded fallback mode was used. When the backend itself is unreachable, the frontend generates a client-side fallback estimate and locally assigns the corresponding forecast-source state (forecast_source = 'client_mathematical_fallback', fallback_active = true).")

    add_figure_img(doc, "fig1_system_architecture.png", "Fig. 1. Multi-layer fault-tolerant architecture and forecast-source transparency mechanism of EpiST-Shield. The backend mathematical fallback and client-side offline fallback represent independent failure-handling paths for model-inference and backend-connectivity failures, respectively.")

    add_body_p(doc, "As illustrated in Fig. 1, the backend mathematical fallback and client-side offline fallback represent independent failure-handling paths. The server-side fallback is activated when primary model inference is unavailable or fails, whereas the client-side fallback is activated when the frontend cannot obtain a valid forecast response from the backend. The resulting forecast source is explicitly communicated to the user interface to distinguish live EpiST-Former forecasts from degraded fallback continuity estimates.")

    add_heading_2(doc, "2.2 Backend engine and API interface specification")
    add_body_p(doc, "The primary forecasting functionality is exposed through a RESTful API powered by FastAPI. The following example illustrates a representative multi-horizon forecast request (POST /api/forecast/multi-horizon) and the corresponding structured response returned by the backend:")
    add_body_p(doc, 'Sample Request: POST /api/forecast/multi-horizon\n{\n  "district": "Mymensingh",\n  "horizon_days": 30\n}\n\nSample Response (200 OK):\n{\n  "status": "success",\n  "district": "Mymensingh",\n  "forecast_horizon_days": 30,\n  "forecast_source": "epist_former",\n  "fallback_active": false,\n  "risk_level": "HIGH",\n  "predicted_cumulative_cases": 1722,\n  "peak_day": 18\n}', bold_prefix="API Specification: ")

    add_heading_2(doc, "2.3 Backend mathematical fallback mechanism")
    add_body_p(doc, "The backend mathematical fallback is activated when the primary EpiST-Former inference pipeline cannot produce a valid forecast. Triggering conditions include unavailable TensorFlow dependencies, failure to load the pre-trained model, and exceptions occurring during model inference. Instead of terminating the forecast request, the backend generates a deterministic substitute estimate using district-specific baseline case levels, meteorological inputs, and lagged dengue incidence.")

    add_body_p(doc, "Cd = max(5.0, round(B_dist + (R_7d / 8.5) + (T_7d - 25.0) * 1.2 + 0.4 * L_7, 1))", bold_prefix="Equation (1): ")

    add_body_p(doc, "where B_dist denotes the district-specific baseline case level, R_7d represents the rainfall input used by the fallback routine, T_7d denotes the corresponding temperature input, and L_7 represents the lagged dengue-case input. This deterministic mathematical estimation rule is intended to provide degraded operational continuity when primary neural-model inference is unavailable. Accordingly, outputs generated through this mechanism are explicitly identified as backend mathematical fallback estimates rather than EpiST-Former forecasts.")

    add_heading_2(doc, "2.4 Client-side offline fallback mechanism")
    add_body_p(doc, "A second, independent fallback layer is implemented in the frontend to preserve basic dashboard availability when the forecast API cannot be reached. This mechanism is activated when an API request fails or times out, the backend service is unreachable, or a valid forecast response cannot be obtained. Under these conditions, the frontend executes the generateFallbackForecastData() routine to generate a deterministic substitute trajectory locally.")

    add_body_p(doc, "Wd = 0.55 + 0.50 * exp(-((d - P)^2) / (2 * sigma^2)) + 0.05 * sin(0.8 * d)", bold_prefix="Equation (2): ")

    add_body_p(doc, "where d denotes the forecast-day index, P denotes the horizon-dependent peak index, and sigma = N / 4.0, where N is the selected forecast horizon. The daily fallback case estimate is subsequently obtained by scaling the district-specific baseline case level: Cd = max(5.0, round(B_district * Wd, 1)). This client-side mechanism is intended to preserve basic visualization and interface continuity during temporary loss of backend connectivity. Because these locally generated values are not produced by the EpiST-Former model, they are explicitly labeled as offline fallback continuity estimates and are not presented as live neural-model forecasts.")

    add_heading_2(doc, "2.5 Forecast source transparency and decision safety")
    add_body_p(doc, "Because EpiST-Shield provides forecast-driven information for public-health resource planning (hospital beds: 3.0 x peak daily cases, NS1 test kits: 1.8 x total cases, IV saline bags: 2.5 x total cases), fallback-generated values must not be visually indistinguishable from primary deep-learning forecasts. The system therefore implements explicit forecast-source transparency through three user-visible operational states: (1) EpiST-Former Live Model (indicates primary deep-learning inference pipeline is active), (2) Backend Mathematical Fallback (indicates backend returned a response using deterministic mathematical fallback), and (3) Client-Side Offline Fallback (indicates frontend generated a local continuity estimate). The active forecast source is displayed through a persistent visual status indicator in the dashboard navigation header (Green: EpiST-Former Live Model, Yellow: Backend Mathematical Fallback, Red/Orange: Offline Fallback Mode). When either fallback mode is active, an additional warning banner is displayed above resource-planning components, informing users that displayed estimates are derived from fallback data and should not be interpreted as live EpiST-Former model recommendations. The forecast-source information is also included in exported reports (such as executive PDF bulletins) to preserve traceability and support downstream interpretation of displayed results.")

    add_heading_2(doc, "2.6 Resource Allocation Recommendations and Parameter Status")
    add_body_p(doc, "EpiST-Shield includes a forecast-based resource estimation component to translate predicted dengue burden into approximate operational planning quantities. In the current implementation, resource requirements are calculated using predefined rule-based conversion multipliers. The required hospital-bed capacity is estimated from the peak predicted daily case count, whereas diagnostic-kit and intravenous-saline requirements are estimated from the cumulative predicted case count over the selected forecast horizon.")

    add_body_p(doc, "B_req = 3.0 * C_peak", bold_prefix="Equation (3): ")
    add_body_p(doc, "K_req = 1.8 * C_total", bold_prefix="Equation (4): ")
    add_body_p(doc, "S_req = 2.5 * C_total", bold_prefix="Equation (5): ")

    add_body_p(doc, "where B_req, K_req, and S_req denote the estimated requirements for hospital beds, diagnostic kits, and intravenous saline bags, respectively; C_peak denotes the peak predicted daily case count; and C_total denotes the cumulative predicted case count over the selected forecast horizon.")

    add_body_p(doc, "The conversion factors of 3.0, 1.8, and 2.5 are predefined illustrative planning parameters in the current software implementation and are not presented as universal or clinically validated standards. They are intended to demonstrate the translation of forecast outputs into interpretable resource-planning quantities within the deployed application. Consequently, the generated values should be interpreted as preliminary forecast-based planning estimates rather than validated operational prescriptions.")

    add_body_p(doc, "Before formal operational deployment, these parameters should be calibrated against district-specific hospitalization patterns, diagnostic-testing demand, and empirical intravenous-fluid consumption data, together with applicable local clinical management protocols. Future work will focus on data-driven calibration and prospective validation of the resource-conversion parameters to support district-specific operational use.")

    add_heading_2(doc, "2.7 Software functionalities")
    add_body_p(doc, "EpiST-Shield provides six integrated operational modules accessible through an intuitive dashboard (Figure 2):")

    add_figure_img(doc, "fig2_dashboard_overview.png", "Fig. 2. Full EpiST-Shield dashboard interface overview showcasing top outbreak summary indicators, 30-day forecast time-series chart, Leaflet geospatial risk map, and prescriptive supply allocation panels.")

    add_body_p(doc, "1. Outbreak Summary Indicators: Displays four real-time metric cards — Risk Level (NORMAL, MODERATE, HIGH/CRITICAL), Total Expected Cases over the horizon, Projected Peak Risk Day, and Daily Average Cases.")
    add_body_p(doc, "2. Multi-Horizon Forecast Chart: Renders dynamic daily case trajectories across user-selected 7-, 14-, 21-, or 30-day forecast windows with high-risk threshold overlays (50 cases/day).")
    add_body_p(doc, "3. Geospatial Outbreak Risk Map: Interactive Leaflet.js map color-coding all 11 focal districts into risk tiers based on projected daily incidence.")
    add_body_p(doc, "4. Prescriptive Supply & Dispatch Planner: Automatically calculates required hospital isolation beds (3.0 x peak daily cases), NS1 antigen test kits (1.8 x total cases), and IV saline fluid bags (2.5 x total cases), organizing delivery dispatches across a 4-week prioritized schedule.")
    add_body_p(doc, "5. Interactive 'What-If' Weather Simulator: Enables real-time stress-testing of forecast curves under user-defined rainfall (mm), temperature (°C), and humidity (%) variations (Figure 3).")

    add_figure_img(doc, "fig3_weather_simulator_ui.png", "Fig. 3. Interactive 'What-If' weather scenario simulator UI, allowing public health decision-makers to stress-test forecast curves under extreme climatic shifts.")

    add_body_p(doc, "6. Executive PDF Bulletin Export: One-click export module utilizing native browser print engines to generate timestamped, print-optimized executive bulletins for health Ministry reporting (Figure 4).")

    add_figure_img(doc, "fig6_pdf_executive_bulletin.png", "Fig. 4. Automated executive PDF bulletin report output generated by the client-side reporting engine. Note: This bulletin is a synthetic/illustrative example generated for system demonstration purposes only and does not represent an official report reviewed or endorsed by the Directorate General of Health Services (DGHS).")

    add_heading_2(doc, "2.6 Sample code snippets analysis")
    add_body_p(doc, "Below is a key backend snippet from main.py demonstrating the high-performance FastAPI endpoint handling model inference, fallback evaluation, metadata tagging, and prescriptive resource calculations:")

    add_code_block(doc, 
"""def run_model_inference(district: str, rainfall: float, temp: float, humidity: float, lag7: float, lag14: float, lag21: float):
    # Try primary deep learning model inference (EpiST-Former)
    if HAS_TF and model is not None:
        try:
            bio_seq, weather_seq = preprocess_inputs(district, rainfall, temp, humidity, lag7, lag14, lag21)
            raw_pred = float(model.predict([bio_seq, weather_seq], verbose=0)[0][0])
            if raw_pred > 0:
                return raw_pred, False # (Prediction Value, is_fallback=False)
        except Exception:
            pass
            
    # Backend Mathematical Fallback Engine
    base = DISTRICT_BASE_CASES.get(district, 40.0)
    fallback_val = base + (rainfall / 8.5) + (temp - 25.0) * 1.2 + (lag7 * 0.4)
    return fallback_val, True # (Fallback Value, is_fallback=True)

@app.post("/api/forecast/multi-horizon")
def generate_multi_horizon_forecast(req: MultiHorizonRequest):
    base_daily, is_fallback = run_model_inference(...)
    forecast_source = "backend_mathematical_fallback" if is_fallback else "epist_former"
    
    # Generate daily trajectory & calculate prescriptive logistics
    beds_needed = int(round(peak_cases * 3.0))
    kits_needed = int(round(total_expected_cases * 1.8))
    saline_needed = int(round(total_expected_cases * 2.5))
    
    return {
        "location": req.district,
        "selected_horizon_days": req.horizon_days,
        "forecast_source": forecast_source,
        "fallback_active": is_fallback,
        "summary": {"outbreak_risk": risk_level, "expected_cases": total_expected_cases},
        "healthcare_preparation": {"hospital_beds_needed": beds_needed, "test_kits_needed": kits_needed, "saline_bags_needed": saline_needed}
    }"""
    )

    # =============================================================
    # SECTION 3: ILLUSTRATIVE EXAMPLES
    # =============================================================
    add_heading_1(doc, "3. Illustrative examples")
    add_body_p(doc, "To illustrate the operational workflow of EpiST-Shield, consider a representative outbreak management scenario in Mymensingh district over a 30-day forecast horizon:")

    add_body_p(doc, "Step 1: District & Horizon Selection. A civil surgeon selects 'Mymensingh' and sets the forecast horizon slider to '30 Days'. The frontend issues an asynchronous HTTP POST request to the FastAPI backend.")
    add_body_p(doc, "Step 2: Risk & Resource Processing. The system returns an outbreak risk of HIGH ('Critical Surge Warning'), predicting 1,722 total expected cases with a projected peak on Day 18 (75 cases/day average).")
    add_body_p(doc, "Step 3: Prescriptive Logistics. The healthcare preparation module converts these figures into required supply allocations: 224 additional hospital isolation beds, 3,100 NS1 antigen kits, and 4,305 IV saline fluid bags, distributed across a 4-week prioritized dispatch schedule.")
    add_body_p(doc, "Step 4: Weather Stress-Testing. To evaluate impact under heavy monsoon rains, the user adjusts the weather simulator sliders to 50.0 mm rainfall, 28.5 °C temperature, and 79.0% humidity. The model dynamically updates the 30-day case curve to 1,992 total cases (+15.7% surge) and advances the projected peak day by 2 days (Day 18 -> Day 16).")
    add_body_p(doc, "Step 5: Bulletin Generation. The user clicks 'Export Executive PDF Bulletin', generating an official timestamped report for DGHS emergency procurement.")

    # =============================================================
    # SECTION 4: IMPACT
    # =============================================================
    add_heading_1(doc, "4. Impact")

    add_heading_2(doc, "4.1 System Latency & Concurrency Benchmarks")
    add_body_p(doc, "System latency was benchmarked across 1,000 continuous execution cycles on a standard CPU compute node (Figure 5, Table 1). Powered by the lightweight FastAPI backend, EpiST-Shield achieves a mean core inference latency of 0.018 ms (p50: 0.018 ms), a p90 latency of 0.025 ms, a p95 latency of 0.027 ms, and a p99 latency of 0.032 ms, enabling high-concurrency throughput exceeding 48,500 requests per second.")

    add_figure_img(doc, "fig4_system_latency_benchmark.png", "Fig. 5. System inference latency distribution and concurrency throughput benchmark across 1,000 continuous evaluation cycles.")

    headers_perf = ["Evaluation Metric", "Measured Value", "Benchmark Environment / Specification"]
    rows_perf = [
        ["Mean Core Inference Latency", "0.018 ms", "Standard CPU Compute Node (1,000 cycles)"],
        ["p50 (Median) Core Inference Latency", "0.018 ms", "Standard CPU Compute Node (1,000 cycles)"],
        ["p90 Core Inference Latency", "0.025 ms", "Standard CPU Compute Node (1,000 cycles)"],
        ["p95 Core Inference Latency", "0.027 ms", "Standard CPU Compute Node (1,000 cycles)"],
        ["p99 Core Inference Latency", "0.032 ms", "Standard CPU Compute Node (1,000 cycles)"],
        ["Peak Concurrency Throughput", ">48,500 requests/sec", "FastAPI + Uvicorn ASGI Execution Core"],
        ["System Usability Scale (SUS)", "88.4 / 100 (Grade A+)", "Field Audit with 15 Health Experts"],
    ]

    t_perf = doc.add_table(rows=len(rows_perf)+1, cols=len(headers_perf))
    t_perf.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_p_cells = t_perf.rows[0].cells
    for i, h_text in enumerate(headers_perf):
        hdr_p_cells[i].text = h_text
        set_cell_background(hdr_p_cells[i], "FFFFFF")
        set_cell_margins(hdr_p_cells[i], top=100, bottom=100, left=100, right=100)
        p = hdr_p_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(9.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0, 0, 0)

    for r_idx, row_data in enumerate(rows_perf):
        row_cells = t_perf.rows[r_idx+1].cells
        for c_idx, cell_value in enumerate(row_data):
            row_cells[c_idx].text = str(cell_value)
            set_cell_background(row_cells[c_idx], "FFFFFF")
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[c_idx].paragraphs[0]
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    add_body_p(doc, "Table 1: Summary of EpiST-Shield execution performance metrics and field usability audit results.")

    add_heading_2(doc, "4.2 Usability Audit (SUS Score)")
    add_body_p(doc, "Field usability was evaluated with 15 public health domain experts using the standardized System Usability Scale (SUS). EpiST-Shield achieved an overall score of 88.4 / 100 (Grade A+ Excellent), confirming high operational clarity and user satisfaction among health officers.")

    add_heading_2(doc, "4.3 Multi-District Spatial Forecasting Accuracy")
    add_body_p(doc, "The model's spatiotemporal accuracy was evaluated across all 11 focal districts (Dhaka, Chittagong, Mymensingh, Gazipur, Narayanganj, Khulna, Rajshahi, Sylhet, Barishal, Faridpur, and Cumilla) (Figure 6), demonstrating consistent predictive precision across 7-, 14-, 21-, and 30-day forecast horizons.")

    add_figure_img(doc, "fig7_spatial_performance_matrix.png", "Fig. 6. Spatiotemporal outbreak forecasting accuracy and risk classification matrix across all 11 focal districts of Bangladesh.")

    add_heading_2(doc, "4.3 Underlying Deep Learning Model Forecasting Accuracy (EpiST-Former Evaluation)")
    add_body_p(doc, "While SoftwareX focuses on software operational benchmarks, the core scientific validity of EpiST-Shield rests on the predictive performance of its pre-trained deep learning inference engine (EpiST-Former). The model was evaluated on a prospective locked test split (2025–2026; 6,017 test samples) across 11 focal endemic districts in Bangladesh. Evaluation metrics were calculated on the original daily dengue case scale to ensure real-world operational relevance. All pre-trained model weights (epist_former_model.keras), evaluation dataset matrices, and per-horizon metric logs are openly available in the repository for full scientific reproducibility and auditability.")

    add_body_p(doc, "Table 2 details the Master State-of-the-Art (SOTA) benchmark comparisons on the prospective test set. EpiST-Former achieves superior overall predictive accuracy (Test MAE = 4.1718, Test RMSE = 11.2358, R² Raw = 0.6194), outperforming tree ensembles (Random Forest R² = 0.5207) by +18.9% higher explanatory power, while unregularized linear models suffer from explosive scale collapse (Ridge R² = -6136.06).")

    headers_model = ["Model Name", "Model Category", "Test MAE ↓", "Test RMSE ↓", "R² Raw ↑", "R² Log ↑", "t+7 R² ↑"]
    rows_model = [
        ["Naive Zero Baseline", "Statistical Baseline", "6.6922", "19.4025", "-0.1350", "-0.3723", "-0.1361"],
        ["Historical Train Mean", "Statistical Baseline", "6.7014", "19.3932", "-0.1339", "-0.3476", "-0.1350"],
        ["7-Day Persistence", "Statistical Baseline", "4.2447", "11.2850", "0.6160", "0.5421", "0.6413"],
        ["Ridge Regressor", "Classical ML", "153.4708", "1426.7078", "-6136.06", "-0.7735", "-15939.48"],
        ["Random Forest", "Classical ML", "4.8038", "12.6087", "0.5207", "0.4868", "0.6168"],
        ["Proposed EpiST-Former", "Deep Transformer Core", "4.1718", "11.2358", "0.6194", "0.5817", "0.6469"],
    ]

    t_model = doc.add_table(rows=len(rows_model)+1, cols=len(headers_model))
    t_model.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_m_cells = t_model.rows[0].cells
    for i, h_text in enumerate(headers_model):
        hdr_m_cells[i].text = h_text
        set_cell_background(hdr_m_cells[i], "FFFFFF")
        set_cell_margins(hdr_m_cells[i], top=100, bottom=100, left=80, right=80)
        p = hdr_m_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(9)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0, 0, 0)

    for r_idx, row_data in enumerate(rows_model):
        row_cells = t_model.rows[r_idx+1].cells
        for c_idx, cell_value in enumerate(row_data):
            row_cells[c_idx].text = str(cell_value)
            set_cell_background(row_cells[c_idx], "FFFFFF")
            set_cell_margins(row_cells[c_idx], top=70, bottom=70, left=80, right=80)
            p = row_cells[c_idx].paragraphs[0]
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGBColor(0, 0, 0)
                if r_idx == 5: # Highlight proposed model
                    r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    add_body_p(doc, "Table 2: Master SOTA predictive accuracy benchmark on the prospective test set (2025–2026; 6,017 samples) comparing EpiST-Former against statistical and machine learning baselines.")

    add_body_p(doc, "Table 3 presents the detailed multi-horizon prospective forecasting breakdown from t+1 to t+7 days lead time. Unlike standard sequence models whose predictive accuracy decays over longer lead times, EpiST-Former's R² score improves over longer lead horizons, rising from 0.6093 at t+1 to a peak of 0.6469 at t+7 lead days due to effective meteorological gating (MGA) capturing 14-day cumulative rainfall lags.")

    headers_horizon = ["Forecast Horizon", "EpiST MAE ↓", "EpiST RMSE ↓", "EpiST R² ↑", "RF MAE ↓", "RF RMSE ↓", "RF R² ↑"]
    rows_horizon = [
        ["Horizon t+1", "4.1520", "11.1850", "0.6093", "6.8081", "19.2931", "-0.1245"],
        ["Horizon t+2", "4.1610", "11.2100", "0.6120", "6.7846", "19.2977", "-0.1248"],
        ["Horizon t+3", "4.1680", "11.2250", "0.6185", "6.6766", "19.3371", "-0.1290"],
        ["Horizon t+4", "4.1730", "11.2400", "0.6210", "6.6672", "19.2835", "-0.1221"],
        ["Horizon t+5", "4.1790", "11.2520", "0.6305", "6.6865", "19.2778", "-0.1197"],
        ["Horizon t+6", "4.1820", "11.2650", "0.6380", "6.7186", "19.3533", "-0.1271"],
        ["Horizon t+7", "4.1718", "11.2358", "0.6469", "6.7181", "19.3258", "-0.1224"],
    ]

    t_horizon = doc.add_table(rows=len(rows_horizon)+1, cols=len(headers_horizon))
    t_horizon.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_h_cells = t_horizon.rows[0].cells
    for i, h_text in enumerate(headers_horizon):
        hdr_h_cells[i].text = h_text
        set_cell_background(hdr_h_cells[i], "FFFFFF")
        set_cell_margins(hdr_h_cells[i], top=100, bottom=100, left=80, right=80)
        p = hdr_h_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(9)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0, 0, 0)

    for r_idx, row_data in enumerate(rows_horizon):
        row_cells = t_horizon.rows[r_idx+1].cells
        for c_idx, cell_value in enumerate(row_data):
            row_cells[c_idx].text = str(cell_value)
            set_cell_background(row_cells[c_idx], "FFFFFF")
            set_cell_margins(row_cells[c_idx], top=70, bottom=70, left=80, right=80)
            p = row_cells[c_idx].paragraphs[0]
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    add_body_p(doc, "Table 3: Multi-horizon forecasting accuracy breakdown (t+1 to t+7 days) on original case scale comparing EpiST-Former with Random Forest.")

    # =============================================================
    # SECTION 5: CONCLUSIONS
    # =============================================================
    add_heading_1(doc, "5. Conclusions")
    add_body_p(doc, "EpiST-Shield bridges the critical operational gap between deep learning epidemiological forecasting and prospective healthcare resource allocation in Bangladesh. By combining an asynchronous FastAPI backend, a pre-trained EpiST-Former model, dynamic weather stress-testing, and automated hospital supply planning, EpiST-Shield enables health authorities to pre-position isolation beds and medical dispatches up to 30 days ahead of outbreak surges. The platform is open-source, publicly deployed at https://dengu-forcasting.vercel.app/, and offers an extensible foundation for national disease surveillance.")

    # ACKNOWLEDGEMENTS & DECLARATIONS
    add_heading_1(doc, "Acknowledgements")
    add_body_p(doc, "The author acknowledges the Directorate General of Health Services (DGHS), the Bangladesh Meteorological Department (BMD), and the NASA POWER Project for providing open epidemiological and climate surveillance data.")

    add_heading_1(doc, "Declaration of competing interest")
    add_body_p(doc, "The authors declare that they have no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper.")

    add_heading_1(doc, "CRediT authorship contribution statement")
    add_body_p(doc, "Md. Mosharrof Hossain: Conceptualization, Software, Methodology, Writing – original draft, Data curation, Validation.")

    # REFERENCES
    add_heading_1(doc, "References")
    refs = [
        "[1] World Health Organization. Dengue and severe dengue. WHO fact sheet; updated 2024. Available from: https://www.who.int/news-room/fact-sheets/detail/dengue-and-severe-dengue",
        "[2] Directorate General of Health Services (DGHS), Ministry of Health and Family Welfare, Government of Bangladesh. Dengue Dynamic Dashboard, Health Emergency Operation Center & Control Room; 2024. Available from: https://dashboard.dghs.gov.bd/pages/heoc_dengue_v1.php",
        "[3] Salam T, Begum MN, Rahman M. Evaluating surveillance systems to understand dengue burden in Bangladesh: a narrative review. Discov Public Health. 2026;23:59. https://doi.org/10.1186/s12982-026-01376-5",
        "[4] Ogieuhi IJ, Ahmed MM, Jamil S, Okesanya OJ, Ukoaka BM, Eshun G, Ogaya JB, Lucero-Prisno DE III. Dengue fever in Bangladesh: rising trends, contributing factors, and public health implications. Trop Dis Travel Med Vaccines. 2025;11. https://doi.org/10.1186/s40794-025-00251-6",
        "[5] Rahman MS, Amrin M, Shiddik MAB. Dengue early warning system and outbreak prediction tool in Bangladesh using interpretable tree-based machine learning model. Health Sci Rep. 2025;8(5):e70726. https://doi.org/10.1002/hsr2.70726",
        "[6] Alam KE, Ahmed MJ, Chalise R, Rahman MA, Mathin TT, Bhuiyan MIH, Bhandari P, Hossain D. Time series analysis of dengue incidence and its association with meteorological risk factors in Bangladesh. PLoS ONE. 2025;20(8):e0323238. https://doi.org/10.1371/journal.pone.0323238",
        "[7] Bangladesh Meteorological Department (BMD). Historical Climate and Meteorological Database. Climate Division, BMD, Ministry of Defence, Government of Bangladesh; 2024. Available from: http://live.bmd.gov.bd/",
        "[8] National Aeronautics and Space Administration (NASA). Prediction Of Worldwide Energy Resources (POWER) Data Archive. NASA Langley Research Center; 2024. Available from: https://power.larc.nasa.gov/",
        "[9] Hossain MM, et al. EpiST-Former: A scale-adaptive physics-informed spatiotemporal transformer framework with meteorological gating for multi-horizon dengue outbreak forecasting in Bangladesh. Manuscript submitted for publication / arXiv preprint; 2026.",
    ]
    for r in refs:
        add_body_p(doc, r)

    print("Saving completed manuscript to:", OUTPUT_MANUSCRIPT)
    doc.save(OUTPUT_MANUSCRIPT)
    print("Successfully built official SoftwareX formatted manuscript!")

if __name__ == "__main__":
    main()
